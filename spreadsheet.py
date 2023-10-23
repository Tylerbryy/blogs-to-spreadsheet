import os
import datetime
import subprocess
import sys


def install_required_modules():
    modules = ["openpyxl", "python-docx", "tkinter", "ttkthemes"]
    for module in modules:
        try:
            __import__(module)
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", module])

install_required_modules()

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from docx import Document
from tkinter import filedialog, messagebox
from ttkthemes import ThemedTk
import tkinter.ttk as ttk

class DOCXExtractor:
    def __init__(self):
        pass

    @staticmethod
    def extract_title_content_from_docx(file_path):
        doc = Document(file_path)
        title = doc.paragraphs[0].text
        content = "\n".join([p.text for p in doc.paragraphs[1:]])
        return title, content

class SpreadsheetPopulator:
    def __init__(self):
        pass

    @staticmethod
    def populate_spreadsheet(directory):
        workbook = Workbook()
        sheet = workbook.active

        # Style for the titles
        bold_font = Font(bold=True, size=14)  # bold font and font size 14
        blue_fill = PatternFill(start_color="6d9eeb", end_color="6d9eeb", fill_type="solid")  # blue fill

        # Add and format titles to the top row
        sheet['A1'] = "Titles"
        sheet['A1'].font = bold_font
        sheet['A1'].fill = blue_fill
        
        sheet['H1'] = "Content"
        sheet['H1'].font = bold_font
        sheet['H1'].fill = blue_fill

        for root, _, files in os.walk(directory):
            for file in files:
                if file.endswith(".docx"):
                    title, content = DOCXExtractor.extract_title_content_from_docx(os.path.join(root, file))
                    next_row = sheet.max_row + 1
                    sheet[f'A{next_row}'] = title
                    sheet[f'H{next_row}'] = content

        timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
        spreadsheet_path = os.path.join(f"output_{timestamp}.xlsx")
        workbook.save(spreadsheet_path)
        return spreadsheet_path


class AppGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("DOCX to Spreadsheet Converter")
        
        # Set the Azure theme
        self.root.set_theme("yaru")

        self.instruction_label = ttk.Label(root, text="Select the root directory containing your DOCX files")
        self.instruction_label.pack(pady=20)

        self.select_button = ttk.Button(root, text="Select Directory", command=self.select_directory)
        self.select_button.pack(pady=20)

    def select_directory(self):
        blog_directory = filedialog.askdirectory(title="Select the root directory containing your DOCX files")
        if not blog_directory:
            return

        try:
            spreadsheet_path = SpreadsheetPopulator.populate_spreadsheet(blog_directory)
            title = "Operation Successful"
            message = f"The extraction process was completed successfully!\n\nSpreadsheet saved at:\n{spreadsheet_path}"
            messagebox.showinfo(title, message)
        except Exception as e:
            messagebox.showerror("Error", str(e))

def main():
    root = ThemedTk()  # Use ThemedTk instead of tk.Tk()
    app = AppGUI(root)
    root.mainloop()

if __name__ == '__main__':
    main()
